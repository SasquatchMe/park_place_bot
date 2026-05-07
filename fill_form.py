"""Fill the Park Place "Material Move-In/Out" form (.docx) with user-supplied values.

The function returns the filled document as bytes (in-memory), so callers
do not need to manage temporary files on disk.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Literal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_PATH = Path(__file__).parent / "template" / "form.docx"

MoveType = Literal["IN", "OUT"]
Category = Literal["Personal Items", "Equipment", "Furniture", "Other"]


@dataclass
class FormValues:
    """Set of values collected from the user before filling the form.

    :param unit: apartment / office number, e.g. ``D1202``.
    :param company: company name in free form.
    :param person_full_name: ФИО ответственного лица.
    :param tel: contact telephone number.
    :param date_str: date in ``DD.MM.YYYY`` format.
    :param move_type: ``IN`` (внос) or ``OUT`` (вынос).
    :param category: one of Personal Items / Equipment / Furniture / Other.
    :param description: free-form items description.
    :param quantity: total number of items being moved.
    :param reason: required only when ``move_type == OUT and quantity > 10``.
    """

    unit: str
    company: str
    person_full_name: str
    tel: str
    date_str: str
    move_type: MoveType
    category: Category
    description: str
    quantity: int
    reason: str = ""


# ---------- low-level helpers ------------------------------------------------


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace paragraph text while preserving the first run's formatting."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    full_text = "".join(run.text for run in paragraph.runs)
    if old not in full_text:
        return False
    _set_paragraph_text(paragraph, full_text.replace(old, new, 1))
    return True


def _cell_replace(cell, old: str, new: str) -> bool:
    for paragraph in cell.paragraphs:
        if _replace_in_paragraph(paragraph, old, new):
            return True
    return False


def _cell_put_centered(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------- fixed coordinates of fields in the template ---------------------

# Coordinates were inferred from the template structure (table 0).
_UNIT_CELL = (1, 0)
_UNIT_PLACEHOLDER = "Unit: ___ - __  __  __  __  "

_COMPANY_CELL = (1, 9)
_COMPANY_PLACEHOLDER = "Company:  _________________________________________"

_PERSON_CELL = (4, 0)
_PERSON_PLACEHOLDER = (
    "Authorized Person’s Name: ___________________________________________________________"
)

_TEL_CELL = (7, 9)
_TEL_PLACEHOLDER = "Tel.: ___________________"

_DATE_CELL = (7, 15)
_DATE_PLACEHOLDER = "Date: ____ / _____ / ____"

_MOVE_IN_CHECKBOX = (16, 1)
_MOVE_OUT_CHECKBOX = (16, 8)

_CATEGORY_CHECKBOXES: dict[Category, tuple[int, int]] = {
    "Personal Items": (18, 1),
    "Equipment": (18, 5),
    "Furniture": (18, 12),
    "Other": (21, 1),
}

_OTHER_LABEL_CELL = (21, 2)
_OTHER_PLACEHOLDER = "Other:___________________________________________"

_BIG_DESCRIPTION_BOX = (25, 4)

# Reason fields (3 underscore lines below the «reason» heading)
_REASON_CELLS = [(28, 1), (29, 1), (30, 1)]
_REASON_LINE_PLACEHOLDER = "____________________________________________________________________________"


# ---------- public API -------------------------------------------------------


def fill_form(values: FormValues) -> bytes:
    """Render the form with the provided values and return the resulting docx bytes.

    :param values: collected form values.
    :return: in-memory docx file content.
    :raises ValueError: when reason is required but missing.
    """
    if values.move_type == "OUT" and values.quantity > 10 and not values.reason.strip():
        raise ValueError("При выносе более 10 предметов обязательно указать причину.")

    document = Document(str(TEMPLATE_PATH))
    table = document.tables[0]

    # Single-line text fields
    _cell_replace(
        _cell(table, _UNIT_CELL),
        _UNIT_PLACEHOLDER,
        f"Unit:  {_format_unit(values.unit)}",
    )
    _cell_replace(
        _cell(table, _COMPANY_CELL),
        _COMPANY_PLACEHOLDER,
        f"Company:  {values.company}",
    )
    _cell_replace(
        _cell(table, _PERSON_CELL),
        _PERSON_PLACEHOLDER,
        f"Authorized Person’s Name:  {values.person_full_name}",
    )
    _cell_replace(
        _cell(table, _TEL_CELL),
        _TEL_PLACEHOLDER,
        f"Tel.:  {values.tel}",
    )
    _cell_replace(
        _cell(table, _DATE_CELL),
        _DATE_PLACEHOLDER,
        f"Date:  {_format_date(values.date_str)}",
    )

    # Move type checkbox
    move_checkbox = _MOVE_IN_CHECKBOX if values.move_type == "IN" else _MOVE_OUT_CHECKBOX
    _cell_put_centered(_cell(table, move_checkbox), "X")

    # Category checkbox
    _cell_put_centered(_cell(table, _CATEGORY_CHECKBOXES[values.category]), "X")

    # Description goes into "Other" line for Other-category, into the big box otherwise.
    if values.category == "Other":
        _cell_replace(
            _cell(table, _OTHER_LABEL_CELL),
            _OTHER_PLACEHOLDER,
            f"Other:  {values.description}",
        )
        # Quantity still goes into the big box
        _put_big_box(table, str(values.quantity))
    else:
        _put_big_box(
            table,
            f"{values.quantity} шт. — {values.description}",
        )

    # Reason (only when OUT and qty > 10)
    if values.move_type == "OUT" and values.quantity > 10 and values.reason.strip():
        _fill_reason(table, values.reason.strip())

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------- helpers used only by the public entry point ---------------------


def _cell(table, coord: tuple[int, int]):
    row, col = coord
    return table.rows[row].cells[col]


def _put_big_box(table, text: str) -> None:
    cell = _cell(table, _BIG_DESCRIPTION_BOX)
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _format_unit(raw: str) -> str:
    """Format unit identifier like ``D1202`` into ``D - 1  2  0  2``.

    Falls back to the raw string when the input does not match the expected pattern.
    """
    cleaned = raw.strip()
    if len(cleaned) >= 2 and cleaned[0].isalpha() and cleaned[1:].isdigit():
        digits = "  ".join(cleaned[1:])
        return f"{cleaned[0].upper()} - {digits}"
    return cleaned


def _format_date(raw: str) -> str:
    """Convert ``DD.MM.YYYY`` (or any ``DD.MM.YY`` variant) into the form layout."""
    cleaned = raw.strip().replace("/", ".").replace("-", ".")
    parts = cleaned.split(".")
    if len(parts) == 3:
        day, month, year = parts
        return f"{day} / {month} / {year}"
    return cleaned


def _fill_reason(table, reason: str) -> None:
    """Distribute the reason text across the three underscore lines under the heading."""
    chunks = _split_reason_into_lines(reason, max_line_chars=70, max_lines=len(_REASON_CELLS))
    for coord, chunk in zip(_REASON_CELLS, chunks):
        _cell_replace(_cell(table, coord), _REASON_LINE_PLACEHOLDER, chunk)


def _split_reason_into_lines(text: str, *, max_line_chars: int, max_lines: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_line_chars:
            current = candidate
            continue
        if current:
            lines.append(current)
        current = word
        if len(lines) == max_lines - 1:
            # last line: dump everything that's left
            remaining = " ".join([current, *words[words.index(word) + 1 :]])
            lines.append(remaining)
            return lines
    if current:
        lines.append(current)
    return lines
